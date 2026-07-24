"""
Resume export: converts a ProfileSchema into an ATS-friendly PDF or DOCX file.
Fulfils PRD 6.1 "AI resume builder (guided Q&A -> formatted resume export as PDF/DOCX)"
and TRD tech stack entry "docx/pdf generation library (e.g., pdf-lib, docx)".
"""
import io
from typing import Any, Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

ACCENT_HEX = "5255F7"


def _contact_line(profile: Dict[str, Any]) -> str:
    parts = [p for p in [profile.get("email"), profile.get("phone"), profile.get("location")] if p]
    return " | ".join(parts)


def build_docx(profile: Dict[str, Any]) -> bytes:
    doc = Document()

    # Name heading
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(profile.get("name") or "Your Name")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x30, 0x2F, 0x94)

    # Contact line
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_p.add_run(_contact_line(profile))
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    def add_section_heading(text: str):
        doc.add_paragraph()
        h = doc.add_paragraph()
        r = h.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x30, 0x2F, 0x94)
        # simple underline rule
        border = doc.add_paragraph()
        border_run = border.add_run("_" * 60)
        border_run.font.size = Pt(6)
        border_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # Skills
    if profile.get("skills"):
        add_section_heading("Skills")
        doc.add_paragraph(", ".join(profile["skills"]))

    # Experience
    if profile.get("experience"):
        add_section_heading("Experience")
        for exp in profile["experience"]:
            line = doc.add_paragraph()
            role_run = line.add_run(f"{exp.get('role', '')} — {exp.get('company', '')}")
            role_run.bold = True
            role_run.font.size = Pt(11)
            if exp.get("duration"):
                dur = line.add_run(f"  ({exp['duration']})")
                dur.italic = True
                dur.font.size = Pt(10)
                dur.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            if exp.get("description"):
                doc.add_paragraph(exp["description"], style="List Bullet")

    # Education
    if profile.get("education"):
        add_section_heading("Education")
        for edu in profile["education"]:
            line = doc.add_paragraph()
            deg_run = line.add_run(f"{edu.get('degree', '')} — {edu.get('school', '')}")
            deg_run.bold = True
            deg_run.font.size = Pt(11)
            if edu.get("year"):
                yr = line.add_run(f"  ({edu['year']})")
                yr.italic = True
                yr.font.size = Pt(10)
                yr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Total experience
    if profile.get("experience_years") is not None:
        add_section_heading("Total Experience")
        doc.add_paragraph(f"{profile['experience_years']} years")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_pdf(profile: Dict[str, Any]) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
        leftMargin=0.7 * inch,
        rightMargin=0.7 * inch,
    )
    styles = getSampleStyleSheet()
    accent = colors.HexColor(f"#{ACCENT_HEX}")

    name_style = ParagraphStyle(
        "NameStyle", parent=styles["Title"], fontSize=22, textColor=accent, spaceAfter=4
    )
    contact_style = ParagraphStyle(
        "ContactStyle", parent=styles["Normal"], fontSize=10, textColor=colors.HexColor("#555555"),
        alignment=1, spaceAfter=14,
    )
    section_style = ParagraphStyle(
        "SectionStyle", parent=styles["Heading2"], fontSize=12, textColor=accent, spaceBefore=14, spaceAfter=4,
    )
    body_style = ParagraphStyle("BodyStyle", parent=styles["Normal"], fontSize=10.5, leading=14)
    subtle_style = ParagraphStyle(
        "SubtleStyle", parent=styles["Normal"], fontSize=9.5, textColor=colors.HexColor("#666666")
    )

    elements = []
    elements.append(Paragraph(profile.get("name") or "Your Name", name_style))
    elements.append(Paragraph(_contact_line(profile), contact_style))
    elements.append(HRFlowable(width="100%", color=colors.HexColor("#DDDDDD"), thickness=0.8))

    if profile.get("skills"):
        elements.append(Paragraph("Skills", section_style))
        elements.append(Paragraph(", ".join(profile["skills"]), body_style))

    if profile.get("experience"):
        elements.append(Paragraph("Experience", section_style))
        for exp in profile["experience"]:
            header = f"<b>{exp.get('role', '')} — {exp.get('company', '')}</b>"
            if exp.get("duration"):
                header += f' <font color="#666666" size="9"><i>({exp["duration"]})</i></font>'
            elements.append(Paragraph(header, body_style))
            if exp.get("description"):
                elements.append(
                    ListFlowable(
                        [ListItem(Paragraph(exp["description"], body_style))],
                        bulletType="bullet",
                        leftIndent=14,
                    )
                )
            elements.append(Spacer(1, 4))

    if profile.get("education"):
        elements.append(Paragraph("Education", section_style))
        for edu in profile["education"]:
            header = f"<b>{edu.get('degree', '')} — {edu.get('school', '')}</b>"
            if edu.get("year"):
                header += f' <font color="#666666" size="9"><i>({edu["year"]})</i></font>'
            elements.append(Paragraph(header, body_style))
            elements.append(Spacer(1, 2))

    if profile.get("experience_years") is not None:
        elements.append(Paragraph("Total Experience", section_style))
        elements.append(Paragraph(f"{profile['experience_years']} years", body_style))

    doc.build(elements)
    return buf.getvalue()
